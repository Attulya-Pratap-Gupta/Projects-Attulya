import webvtt
from docx import Document
#package to work with a .vtt file
start_list,name_list,speech_list, spk_num_lst = [],[],[],[]
#Lists for storing start timing index, name of the speaker, and content

for caption in webvtt.read('Team23_220808.vtt'):  #Change .vtt file name
    #Reading through webvtt file
    start_list.append(caption.start[:8])
    #Fetching the start time and appending to list
    #end = caption.end Fetches the end time
    full_text = caption.text
    name_end_index = full_text.find(':')
    
    if name_end_index == -1:
        #Condition to exclude null values
        name_list.append(' ')
        speech_list.append(full_text.capitalize())
        continue
    
    name_list.append(full_text[:name_end_index])
    #Fetching the speaker name
    speech = full_text[name_end_index+2:]
    speech_list.append(speech.capitalize())
    #Getting the words spoken by the speaker

identities = set()
counter=0
for value in name_list:
    new_value = ''
    for char in value:
        if char != ' ':
            new_value += char
    identities.add(new_value)
    name_list[counter] = new_value
    counter+=1
identities = sorted(list(identities))
for value in name_list:
    spk_num_lst.append(identities.index(value))

word_document = Document()
audio_file = word_document.add_heading('Audio file',level=1)
audio_file = word_document.add_paragraph('Team23_220808')
#Change team name and link according to the required video
transcript = word_document.add_heading('Transcript',level=1)
for i in range(len(name_list)):
    transcript = word_document.add_paragraph(str(start_list[i])+' Speaker '+str(spk_num_lst[i]))
    transcript = word_document.add_paragraph(str(speech_list[i]))
word_document.save('test1.docx')
#Change File name accordingly
    








    
    
